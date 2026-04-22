"""
Generate IBS Project Report for Team 22 in DOCX format,
matching the structure of the Team-21 report.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

FONT_NAME = 'Cambria Math'

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:color'): 'auto',
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = FONT_NAME
    return h

def add_para(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=6, justified=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if alignment is not None:
        p.alignment = alignment
    elif justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_centered(doc, text, bold=False, size=14):
    return add_para(doc, text, bold=bold, size=size, alignment=WD_ALIGN_PARAGRAPH.CENTER, justified=False)

def add_image_if_exists(doc, path, width=Inches(5.0), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.space_before = Pt(12)
        last_paragraph.paragraph_format.space_after = Pt(6)
        if caption:
            add_para(doc, caption, italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, justified=False)
        return True
    else:
        add_para(doc, f"[Image not found: {os.path.basename(path)}]", italic=True, size=10)
        return False

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME

    return table

def generate_report():
    doc = Document()

    # --- Set page margins (1 inch = 2.54 cm) ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)   # standard left margin
        section.right_margin = Cm(3.18)  # standard right margin

    # --- Set default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Set heading fonts ---
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = FONT_NAME
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

    # --- Set List styles font ---
    for style_name in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[style_name]
            ls.font.name = FONT_NAME
            ls.font.size = Pt(12)
        except KeyError:
            pass

    # ========== TITLE PAGE ==========
    for _ in range(3):
        doc.add_paragraph()

    add_centered(doc, 'Structure-Guided Graph Neural Network for Predicting Protein Melting Temperature from Sequence and 3D Features', bold=True, size=16)
    doc.add_paragraph()
    add_centered(doc, 'PROJECT REPORT', bold=True, size=14)
    doc.add_paragraph()
    add_centered(doc, 'Submitted to', size=12)
    add_centered(doc, 'Amrita School of Computing, Amrita Vishwa Vidyapeetham, Chennai', size=12)
    doc.add_paragraph()
    add_centered(doc, 'in partial fulfilment for the award of the degree of', size=12)
    add_centered(doc, 'BACHELOR OF TECHNOLOGY', bold=True, size=12)
    add_centered(doc, 'COMPUTER SCIENCE ENGINEERING - ARTIFICIAL INTELLIGENCE', bold=True, size=12)
    doc.add_paragraph()
    add_centered(doc, 'By', size=12)
    add_para(doc, '    Achintya TR \u2013 CH.SC.U4AIE24003', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '    Lishanth SS \u2013 CH.SC.U4AIE24083', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    h = doc.add_heading('Faculty Mentor:   DR I. R. OVIYA', level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_centered(doc, 'AMRITA SCHOOL OF COMPUTING,', size=12)
    add_centered(doc, 'AMRITA VISHWA VIDYAPEETHAM, CHENNAI CAMPUS', size=12)
    doc.add_paragraph()
    add_centered(doc, 'March 2026', size=12)

    doc.add_page_break()

    # ========== BONAFIDE CERTIFICATE ==========
    add_centered(doc, 'BONAFIDE CERTIFICATE', bold=True, size=14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, (
        'It is declared that this project entitled \u201cStructure-Guided Graph Neural Network for Predicting '
        'Protein Melting Temperature from Sequence and 3D Features\u201d by Achintya TR, Lishanth SS was undertaken '
        'under my direct supervision for End semester evaluation for the course 22BIO211- Intelligence of '
        'Biological systems - 2 in the even semester of the Academic Year 2025-2026.'
    ), size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, (
        'Dr. I. R. Oviya\nAsst. Professor\nDept. of CSE- Artificial Intelligence\n'
        'Amrita School of Computing\nAmrita Vishwa Vidyapeetham, Chennai Campus'
    ), size=12)

    doc.add_page_break()

    # ========== DECLARATION ==========
    add_centered(doc, 'DECLARATION', bold=True, size=14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, (
        'We solemnly declare that the project titled \u201cStructure-Guided Graph Neural Network for Predicting '
        'Protein Melting Temperature from Sequence and 3D Features\u201d, submitted for the End Semester Project '
        'evaluation during the even semester of the academic year 2025-2026, for the course 22BIO211 - '
        'Intelligence of Biological Systems- 2, offered by the Department of Computer Science Engineering - '
        'Artificial Intelligence, is a true representation of our original work. This project was undertaken '
        'under the guidance of Dr. I. R. Oviya, Assistant Professor, Department of Computer Science Engineering '
        '- Artificial Intelligence, Amrita School of Computing, Amrita Vishwa Vidyapeetham, Chennai Campus. '
        'We further affirm that the contents of this project have not been submitted nor will be submitted to '
        'any other institution or entity for the fulfillment of any degree or qualification.'
    ), size=12)

    doc.add_page_break()

    # ========== ACKNOWLEDGEMENTS ==========
    add_centered(doc, 'ACKNOWLEDGEMENTS', bold=True, size=14)
    doc.add_paragraph()
    add_para(doc, (
        'We extend our sincere gratitude to all those who have supported and inspired us throughout the '
        'duration of this project on Structure-Guided Graph Neural Network for Predicting Protein Melting '
        'Temperature from Sequence and 3D Features.'
    ))
    add_para(doc, (
        'First and foremost, we would like to thank our Course Faculty, Dr. I. R. Oviya, whose expertise, '
        'understanding, and patience, added considerably to our research experience. We are immensely grateful '
        'for her guidance in formulating the research question and framework.'
    ))
    add_para(doc, (
        'We also wish to express our deep appreciation to the members of our team, Achintya TR and Lishanth SS, '
        'for their collaborative effort, insightful discussions, and encouragement, which were invaluable '
        'throughout the research process.'
    ))
    add_para(doc, (
        'We are also thankful to the TemStaPro database, the AlphaFold Protein Structure Database, and the '
        'ESM-2 protein language model team, and all the authors and researchers whose publications and databases '
        'provided the essential data and insights that formed the backbone of this study.'
    ))
    add_para(doc, (
        'Gratitude is also due to our family and friends, who have provided us with endless encouragement '
        'and were always there to lift our spirits during the challenging times of this research.'
    ))
    add_para(doc, (
        'This project not only represents our academic endeavours but also the support and belief of '
        'everyone mentioned above. Thank you.'
    ))

    doc.add_page_break()

    # ========== CONTRIBUTION STATEMENT ==========
    add_centered(doc, 'CONTRIBUTION STATEMENT', bold=True, size=14)
    doc.add_paragraph()
    add_para(doc, (
        'The undersigned members of Team 22 hereby submit the contribution statement for the project titled '
        '"Structure-Guided Graph Neural Network for Predicting Protein Melting Temperature from Sequence and '
        '3D Features". This document outlines the individual contributions of each team member towards the '
        'completion of the project.'
    ))
    add_para(doc, 'Team Number: 22', bold=True)
    doc.add_paragraph()

    make_table(doc,
        ['No.', 'Name', 'Register Number', 'Percentage Contribution', 'Contribution', 'Signature'],
        [
            ['1.', 'Achintya TR', 'CH.SC.U4AIE24003', '50%',
             'Research, 50% pipeline building, 50% report writing', ''],
            ['2.', 'Lishanth SS', 'CH.SC.U4AIE24083', '50%',
             '50% pipeline building, 50% report writing', ''],
        ]
    )

    doc.add_paragraph()
    add_para(doc, (
        'This statement is presented to attest to the distribution of work among the team members, ensuring '
        'that each member\'s contributions are duly acknowledged. Each member has agreed to the accuracy of '
        'this statement and has committed to uphold the integrity of the collaborative work presented.'
    ))
    doc.add_paragraph()
    add_para(doc, 'Date: 30/03/2026\n')
    add_para(doc, 'Place: Chennai')

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_centered(doc, 'TABLE OF CONTENTS', bold=True, size=14)
    doc.add_paragraph()

    make_table(doc,
        ['CHAPTER NO.', 'TITLE', 'PAGE NO.'],
        [
            ['1', 'INTRODUCTION', '9'],
            ['', '1.1 Significance of Protein Melting Temperature Prediction', '9'],
            ['', '1.1.1 Thermal Stability as a Bottleneck in Protein Engineering', '9'],
            ['', '1.2 Protein Stability in Pharmaceutical and Industrial Context', '10'],
            ['', '1.3 Graph Neural Networks for Protein Structure Learning', '11'],
            ['2', 'LITERATURE REVIEW', '13'],
            ['3', 'METHODOLOGY', '16'],
            ['', '3.1 Data Preprocessing', '16'],
            ['', '3.1.1 Description of Dataset', '16'],
            ['', '3.1.2 Encoding of Features', '17'],
            ['', '3.2 Graph Construction from 3D Structures', '19'],
            ['', '3.3 Model Architecture', '20'],
            ['', '3.3.1 GCN Model', '20'],
            ['', '3.3.2 GIN Model', '21'],
            ['', '3.4 Model Training', '22'],
            ['', '3.4.1 Train/Test Split', '22'],
            ['', '3.4.2 Optimizer and Loss Function', '22'],
            ['', '3.4.3 Training Dynamics and Early Stopping', '23'],
            ['', '3.4.4 Evaluation Metrics', '23'],
            ['4', 'RESULTS AND DISCUSSION', '24'],
            ['', '4.1 Training Performance', '24'],
            ['', '4.2 Model Comparison', '25'],
            ['', '4.3 Prediction Scatter Plots', '26'],
            ['', '4.4 Discussion', '27'],
            ['5', 'CONCLUSION', '30'],
            ['6', 'REFERENCES', '31'],
        ]
    )

    doc.add_page_break()

    # ========== ABSTRACT ==========
    add_centered(doc, 'ABSTRACT', bold=True, size=14)
    doc.add_paragraph()
    add_para(doc, (
        'Protein melting temperature (Tm) is a fundamental biophysical property that determines the thermal '
        'stability, shelf-life, and functional viability of proteins in therapeutic, industrial, and research '
        'applications. Accurate computational prediction of Tm can accelerate protein engineering by reducing '
        'the need for costly and time-consuming experimental assays such as differential scanning calorimetry. '
        'This study presents a structure-guided Graph Neural Network (GNN) approach for predicting protein '
        'melting temperature directly from amino acid sequences and three-dimensional structural information. '
        'The pipeline integrates ESM-2 protein language model embeddings (320-dimensional per-residue features) '
        'with residue-level contact graphs constructed from AlphaFold-predicted 3D structures using an 8\u00c5 '
        'distance threshold. Two Graph Neural Network architectures\u2014Graph Convolutional Network (GCN) and '
        'Graph Isomorphism Network (GIN)\u2014were implemented and evaluated on a large-scale dataset of 100,000 '
        'proteins derived from the TemStaPro database, which contains 943,605 experimentally annotated protein '
        'sequences with Tm values ranging from 4\u00b0C to 100\u00b0C. The GIN model achieved a Pearson correlation '
        'coefficient of 0.745, outperforming the GCN model (r = 0.720), demonstrating that the more expressive '
        'GIN architecture better captures the structural determinants of thermal stability. Training employed '
        'MSE loss with Adam optimizer, L2 regularization, and early stopping. These results demonstrate that '
        'combining protein language model embeddings with structure-based graph representations provides a '
        'powerful framework for absolute Tm prediction across diverse protein families, offering a scalable '
        'alternative to sequence-only or structure-only approaches.'
    ))

    doc.add_page_break()

    # ========== CHAPTER 1: INTRODUCTION ==========
    add_heading_styled(doc, 'CHAPTER 1', level=1)
    add_heading_styled(doc, 'INTRODUCTION', level=1)

    add_heading_styled(doc, '1.1 Significance of Protein Melting Temperature Prediction', level=2)
    add_heading_styled(doc, '1.1.1 Thermal Stability as a Bottleneck in Protein Engineering', level=3)
    add_para(doc, (
        'Protein engineers design novel sequences for therapeutics and industrial applications, yet ensuring '
        'sufficient thermal stability (high melting temperature, Tm) remains a critical bottleneck. Experimental '
        'techniques such as differential scanning calorimetry (DSC) and differential scanning fluorimetry (DSF) '
        'are costly, time-consuming, and unsuitable for screening vast sequence spaces. The melting temperature '
        'of a protein, defined as the temperature at which 50% of the protein population is in the unfolded '
        'state, is a key indicator of thermodynamic stability and directly impacts formulation, storage, and '
        'functional performance of biopharmaceuticals.'
    ))
    add_para(doc, (
        'According to recent estimates, late-stage failures due to poor stability account for a significant '
        'fraction of drug development attrition. The ability to computationally predict Tm before synthesis '
        'and expression would dramatically reduce experimental screening costs and accelerate the design cycle '
        'for therapeutic antibodies, industrial enzymes, and engineered proteins. This motivates the development '
        'of data-driven machine learning approaches that can learn the complex relationship between protein '
        'sequence, structure, and thermal stability.'
    ))

    add_heading_styled(doc, '1.2 Protein Stability in Pharmaceutical and Industrial Context', level=2)
    add_para(doc, (
        'In the pharmaceutical industry, protein stability is a critical quality attribute (CQA) that affects '
        'drug safety, efficacy, and manufacturability. Regulatory agencies require thorough characterization '
        'of protein stability during drug development. Industrial enzymes used in detergents, food processing, '
        'and biofuel production must operate at elevated temperatures, making thermostability a primary design '
        'objective. The growing interest in extremophilic enzymes and de novo protein design has further '
        'amplified the need for accurate Tm prediction tools that work across diverse protein families and folds.'
    ))
    add_para(doc, (
        'Current computational approaches face several limitations. Methods based on sequence features alone '
        '(amino acid composition, dipeptide frequencies) cannot capture long-range spatial interactions that '
        'stabilize the folded state. Structure-based approaches using molecular dynamics simulations are '
        'computationally prohibitive for large-scale screening. More recent machine learning methods that '
        'predict changes in stability upon mutation (\u0394\u0394G) require a reference protein and do not '
        'directly predict absolute Tm values. This gap motivates the development of structure-aware learning '
        'methods that can predict absolute melting temperatures directly.'
    ))

    add_heading_styled(doc, '1.3 Graph Neural Networks for Protein Structure Learning', level=2)
    add_para(doc, (
        'Graph Neural Networks (GNNs) provide a natural framework for learning from protein structures. By '
        'representing proteins as graphs\u2014where nodes correspond to amino acid residues and edges represent '
        'spatial contacts between residues\u2014GNNs can learn hierarchical structural features through iterative '
        'message passing. This approach captures both local secondary structure elements and long-range tertiary '
        'contacts that are critical for stability.'
    ))
    add_para(doc, (
        'Recent advances in protein language models, particularly ESM-2 from Meta AI, have demonstrated that '
        'pre-trained transformer models can generate rich, context-aware embeddings for each residue in a protein '
        'sequence. These embeddings encode evolutionary, structural, and functional information learned from '
        'millions of protein sequences. By combining ESM-2 embeddings as node features with structure-derived '
        'contact graphs, we hypothesize that a hybrid GNN model can capture both sequence-level and '
        'structure-level determinants of thermal stability.'
    ))
    add_para(doc, (
        'This project addresses the challenge of absolute Tm prediction by framing it as a regression task '
        'using both sequence features (ESM-2 embeddings) and 3D protein graphs (AlphaFold structures). Two '
        'GNN architectures\u2014GCN and GIN\u2014are implemented, trained on 100,000 proteins from the TemStaPro '
        'database, and systematically evaluated to quantify the contribution of structural information to '
        'Tm prediction accuracy.'
    ))

    doc.add_page_break()

    # ========== CHAPTER 2: LITERATURE REVIEW ==========
    add_heading_styled(doc, 'CHAPTER 2', level=1)
    add_heading_styled(doc, 'LITERATURE REVIEW', level=1)

    add_heading_styled(doc, '2.1 Traditional Approaches to Protein Stability Prediction', level=2)
    add_para(doc, (
        'Early computational approaches to protein stability prediction relied on empirical potential functions '
        'and statistical analysis of known protein structures. Methods such as FoldX and Rosetta used '
        'physics-based energy functions to estimate changes in folding free energy (\u0394\u0394G) upon mutation. '
        'While effective for point mutations in well-characterized protein families, these methods are '
        'computationally expensive and struggle to generalize across diverse protein architectures.'
    ))
    add_para(doc, (
        'Machine learning approaches using hand-crafted sequence features\u2014such as amino acid composition, '
        'dipeptide frequencies, and physicochemical property averages\u2014have achieved moderate success in '
        'classifying proteins as thermophilic versus mesophilic. However, these methods fundamentally cannot '
        'capture the complex spatial interactions that govern protein stability, limiting their accuracy for '
        'continuous Tm prediction across the full temperature range.'
    ))

    add_heading_styled(doc, '2.2 Deep Learning and Protein Language Models', level=2)
    add_para(doc, (
        'The advent of large-scale protein language models has transformed computational biology. Models such as '
        'ESM-2 (Lin et al., 2023), ProtTrans (Elnaggar et al., 2022), and ProtT5 learn rich per-residue '
        'representations from evolutionary patterns in millions of protein sequences. These embeddings capture '
        'structural, functional, and evolutionary information without requiring explicit 3D structures.'
    ))
    add_para(doc, (
        'TemBERTure (2024) advanced protein thermostability prediction using a BERT-like protein language model, '
        'achieving direct Tm prediction from sequences alone. The TemBERTure-Tm model demonstrated that PLM '
        'embeddings encode meaningful information about thermal stability. DeepTM (2023) similarly employed '
        'sequence-based deep learning for Tm prediction of thermophilic proteins, incorporating amino acid '
        'composition and evolutionary features. However, both approaches are limited by their inability to '
        'capture 3D structural contacts that play a critical role in protein stability.'
    ))

    add_heading_styled(doc, '2.3 Graph Neural Networks for Molecular and Protein Property Prediction', level=2)
    add_para(doc, (
        'Graph Neural Networks have emerged as powerful tools for learning from molecular and protein structures. '
        'Kipf and Welling (2017) introduced Graph Convolutional Networks (GCNs) for semi-supervised learning on '
        'graphs, which have been widely adopted for molecular property prediction. Xu et al. (2019) proposed the '
        'Graph Isomorphism Network (GIN), proving that it achieves maximal expressiveness among message-passing '
        'GNNs by being as powerful as the Weisfeiler-Lehman graph isomorphism test.'
    ))
    add_para(doc, (
        'In the protein stability domain, ThermoAGT-GA (2024) employed restricted-attention subgraph GNNs on '
        'residue contact maps for predicting \u0394\u0394G upon mutation. ProSTAGE (2024) fused ProtT5 embeddings '
        'with GCN-processed structural features for \u0394\u0394G prediction, demonstrating strong cross-dataset '
        'generalization. Stability Oracle (2024) used atom-level graph-transformers with distance-biased '
        'attention for identifying thermodynamically stable sequences. However, all these approaches focus on '
        '\u0394\u0394G prediction for mutations rather than absolute Tm regression for arbitrary proteins.'
    ))

    add_heading_styled(doc, '2.4 Research Gap and Contribution of This Work', level=2)
    add_para(doc, 'Based on the literature review, the following key research gaps were identified:')
    gaps = [
        'Most GNN-based approaches predict \u0394\u0394G for mutations, not absolute Tm for arbitrary sequences.',
        'Sequence-only models (TemBERTure, DeepTM) ignore 3D structural effects critical for stability.',
        'No existing work combines PLM embeddings with structure-based GNNs specifically for Tm regression.',
        'Large-scale Tm datasets (e.g., TemStaPro with 943K proteins) remain underutilized for GNN-based prediction.',
        'Systematic comparison of GNN architectures (GCN vs GIN) on Tm prediction is lacking.',
    ]
    for gap in gaps:
        p = doc.add_paragraph(gap, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT_NAME

    add_para(doc, (
        'This work fills these gaps by: (1) presenting a complete pipeline from raw FASTA sequences to '
        'AlphaFold structure retrieval to GNN-based Tm regression; (2) combining ESM-2 protein language model '
        'embeddings with residue-contact graphs; (3) systematically comparing GCN and GIN architectures on a '
        '100,000-protein subset; and (4) demonstrating that structure-guided GNNs achieve Pearson correlations '
        'above 0.7 for absolute Tm prediction.'
    ))

    doc.add_page_break()

    # ========== CHAPTER 3: METHODOLOGY ==========
    add_heading_styled(doc, 'CHAPTER 3', level=1)
    add_heading_styled(doc, 'METHODOLOGY', level=1)

    add_heading_styled(doc, '3.1 Data Preprocessing', level=2)
    add_heading_styled(doc, '3.1.1 Description of Dataset', level=3)
    add_para(doc, (
        'The dataset for this study was sourced from the TemStaPro database (Bioinformatics, 2024), which '
        'provides experimentally reported melting temperature (Tm) annotations for proteins. The raw data was '
        'obtained from the file TemStaPro-Major-30-imbal-training.fasta, containing protein sequences in FASTA '
        'format with headers encoding the organism ID, UniProt accession, and experimental Tm value.'
    ))
    add_para(doc, (
        'The FASTA file was parsed using a custom Python script (create_database.py) that extracted 943,605 '
        'protein entries with three fields: UniProt ID, amino acid sequence, and melting temperature (Tm in \u00b0C). '
        'The Tm values in the dataset range from 4\u00b0C to 100\u00b0C, with a mean of 37.5\u00b0C and a standard '
        'deviation of 14.3\u00b0C. The distribution is right-skewed, with the majority of proteins having Tm '
        'values between 25\u00b0C and 45\u00b0C.'
    ))

    # Dataset statistics table
    make_table(doc,
        ['Statistic', 'Value'],
        [
            ['Total Records', '943,605'],
            ['Mean Tm', '37.53\u00b0C'],
            ['Std Dev', '14.25\u00b0C'],
            ['Min Tm', '4.00\u00b0C'],
            ['25th Percentile', '28.00\u00b0C'],
            ['Median Tm', '30.00\u00b0C'],
            ['75th Percentile', '45.00\u00b0C'],
            ['Max Tm', '100.00\u00b0C'],
        ]
    )
    add_para(doc, 'Table 1: Summary statistics of the TemStaPro dataset.', italic=True, size=10,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, (
        'For computational feasibility, a 100,000-protein subset (data_100k.csv) was randomly sampled from '
        'the full dataset using a stratified sampling approach to maintain the temperature distribution. '
        'Additionally, smaller development subsets (500, 5,000, and 20,000 samples) were created using '
        'create_dev_subset.py for rapid prototyping and pipeline validation.'
    ))

    add_heading_styled(doc, '3.1.2 Structure Acquisition from AlphaFold', level=3)
    add_para(doc, (
        'Since the GNN approach requires 3D protein structures, AlphaFold-predicted structures were downloaded '
        'for each protein in the dataset. A custom download script (download_pdb.py) was developed to automate '
        'this process. The script queries the AlphaFold Database API (https://alphafold.ebi.ac.uk/api/prediction/) '
        'using UniProt accession IDs, resolves the latest model version, and downloads PDB files. Multi-threaded '
        'downloading with 10 concurrent workers was employed to accelerate the process. The script includes '
        'resume capability, skipping proteins whose structures have already been downloaded.'
    ))

    add_heading_styled(doc, '3.1.3 Encoding of Features: ESM-2 Protein Language Model Embeddings', level=3)
    add_para(doc, (
        'Node features for each residue in the protein graph are generated using ESM-2 (Evolutionary Scale '
        'Modeling), a state-of-the-art protein language model developed by Meta AI. Specifically, we used the '
        'esm2_t6_8M_UR50D model (8 million parameters, 6 transformer layers), which outputs 320-dimensional '
        'per-residue embeddings. This compact model was chosen for computational efficiency while retaining '
        'rich evolutionary and structural information learned from the UniRef50 database.'
    ))
    add_para(doc, (
        'The embedding extraction process works as follows: (1) Each protein sequence is tokenized using the '
        'ESM alphabet; (2) The tokenized sequence is passed through the ESM-2 model to obtain representations '
        'from the final transformer layer (layer 6); (3) The special <cls> and <eos> tokens are removed, '
        'yielding a per-residue embedding matrix of shape (L \u00d7 320), where L is the sequence length. '
        'Sequences longer than 1,022 residues are truncated to prevent GPU memory overflow. The ESM-2 model '
        'is run in evaluation mode with gradient computation disabled for efficiency.'
    ))

    add_heading_styled(doc, '3.2 Graph Construction from 3D Structures', level=2)
    add_para(doc, (
        'Protein graphs are constructed from AlphaFold-predicted PDB structures using a distance-based '
        'approach implemented in the ProteinGraphDataset class (gnn_data.py). The graph construction follows '
        'these steps:'
    ))
    steps = [
        'C-alpha atoms are extracted from the PDB file using BioPandas, providing the 3D coordinates for each residue.',
        'A pairwise Euclidean distance matrix is computed between all C-alpha atoms.',
        'An adjacency matrix is created by thresholding: two residues are connected by an edge if their C-alpha distance is less than 8.0 \u00c5 (Angstroms) and greater than 0 (excluding self-loops).',
        'The resulting sparse adjacency is converted to a PyTorch Geometric edge_index tensor.',
        'For proteins where the structure length does not match the sequence length, or where the PDB file is missing, a fallback linear chain graph is constructed (each residue connected to its sequential neighbors).',
    ]
    for step in steps:
        p = doc.add_paragraph(step, style='List Number')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT_NAME

    add_para(doc, (
        'The 8\u00c5 threshold is a standard choice in structural bioinformatics, capturing both direct contacts '
        'and near-contacts that contribute to the stability of the protein fold. This graph representation '
        'enables the GNN to learn from long-range spatial interactions that sequential models cannot capture.'
    ))

    add_para(doc, (
        'Each protein is represented as a PyTorch Geometric Data object containing: (1) x: node feature matrix '
        'of shape (L \u00d7 320) from ESM-2 embeddings; (2) edge_index: edge connectivity in COO format; '
        '(3) y: target melting temperature in \u00b0C; (4) uniprot_id: protein identifier for traceability.'
    ))

    add_heading_styled(doc, '3.3 Model Architecture', level=2)
    add_heading_styled(doc, '3.3.1 GCN Model (Graph Convolutional Network)', level=3)
    add_para(doc, (
        'The GCN model architecture (GNNModel class in gnn_model.py) consists of three Graph Convolutional '
        'layers with ReLU activation, followed by global mean pooling and a linear regression head. The '
        'architecture is summarized as follows:'
    ))

    make_table(doc,
        ['Layer', 'Input Dim', 'Output Dim', 'Activation'],
        [
            ['GCNConv 1', '320', '64', 'ReLU'],
            ['GCNConv 2', '64', '64', 'ReLU'],
            ['GCNConv 3', '64', '64', '\u2014'],
            ['Global Mean Pool', '64', '64', '\u2014'],
            ['Dropout', '64', '64', 'p=0.5'],
            ['Linear', '64', '1', '\u2014'],
        ]
    )
    add_para(doc, 'Table 2: GCN model architecture.', italic=True, size=10,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, (
        'The GCN layers perform spectral-domain graph convolution as introduced by Kipf and Welling (2017). '
        'Each layer aggregates feature information from neighboring nodes using a normalized adjacency matrix, '
        'enabling the model to learn hierarchical representations of the protein graph. Global mean pooling '
        'aggregates node-level features into a fixed-size graph-level representation, which is then mapped '
        'to a single Tm prediction through a linear layer with dropout regularization.'
    ))

    add_heading_styled(doc, '3.3.2 GIN Model (Graph Isomorphism Network)', level=3)
    add_para(doc, (
        'The GIN model (GINModel class in gnn_model.py) employs a more expressive architecture based on the '
        'Graph Isomorphism Network framework by Xu et al. (2019). GIN achieves maximal discriminative power '
        'among message-passing GNNs by using learnable MLPs in the aggregation step. The architecture consists '
        'of two GIN convolution blocks with batch normalization, followed by global sum pooling and a two-layer '
        'classifier.'
    ))

    make_table(doc,
        ['Layer', 'Input Dim', 'Output Dim', 'Details'],
        [
            ['GINConv 1 (MLP)', '320', '128', '2-layer MLP + BatchNorm + ReLU'],
            ['BatchNorm 1', '128', '128', '\u2014'],
            ['GINConv 2 (MLP)', '128', '128', '2-layer MLP + BatchNorm + ReLU'],
            ['BatchNorm 2', '128', '128', '\u2014'],
            ['Global Sum Pool', '128', '128', '\u2014'],
            ['Linear 1 + ReLU', '128', '128', '\u2014'],
            ['Dropout', '128', '128', 'p=0.5'],
            ['Linear 2', '128', '1', '\u2014'],
        ]
    )
    add_para(doc, 'Table 3: GIN model architecture.', italic=True, size=10,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, (
        'Key differences from the GCN model include: (1) GIN uses learnable MLPs within each convolution layer '
        'instead of fixed aggregation, providing greater expressiveness; (2) A hidden dimension of 128 (vs 64 '
        'in GCN) increases model capacity; (3) Global sum pooling (vs mean pooling) preserves information about '
        'graph size, which may correlate with protein molecular weight and stability; (4) Batch normalization '
        'after each convolution block stabilizes training.'
    ))

    add_heading_styled(doc, '3.3.3 Model Parameter Summary', level=3)
    add_para(doc, (
        'Both models are relatively compact. The GCN model has approximately 21,500 parameters, and the GIN '
        'model has approximately 126,000 parameters due to the larger hidden dimension and MLP-based convolutions. '
        'This moderate model size helps prevent overfitting, especially when training on the 100,000-sample '
        'subset with diverse protein families.'
    ))

    add_heading_styled(doc, '3.4 Model Training', level=2)
    add_heading_styled(doc, '3.4.1 Train/Test Split', level=3)
    add_para(doc, (
        'The 100,000-protein dataset was split into 80% training (80,000 samples) and 20% testing (20,000 '
        'samples) using random permutation. The split was performed at the data loading stage to ensure '
        'reproducibility. Both splits were loaded using PyTorch Geometric DataLoader with a batch size of 32.'
    ))

    add_heading_styled(doc, '3.4.2 Optimizer and Loss Function', level=3)
    add_para(doc, (
        'The model was optimized using the Adam optimizer with a learning rate of 0.001 and L2 weight decay '
        'regularization (\u03bb = 1\u00d710\u207b\u2074). The loss function was Mean Squared Error (MSE), '
        'which is the standard choice for regression tasks. MSE penalizes large prediction errors quadratically, '
        'encouraging the model to minimize the variance of its predictions.'
    ))

    add_heading_styled(doc, '3.4.3 Training Dynamics and Early Stopping', level=3)
    add_para(doc, (
        'Training was conducted for a maximum of 50 epochs with early stopping based on test MSE. The patience '
        'parameter was set to 8 epochs: if the test loss did not improve for 8 consecutive epochs, training was '
        'terminated and the best model (lowest test MSE) was saved. This strategy prevents overfitting by '
        'stopping training when the model begins to memorize the training data rather than learning generalizable '
        'patterns. In practice, early stopping was triggered around epoch 24 for the GIN model.'
    ))

    add_heading_styled(doc, '3.4.4 Evaluation Metrics', level=3)
    add_para(doc, 'The following regression metrics were used to evaluate model performance:')
    metrics = [
        'Mean Squared Error (MSE): Average squared difference between predicted and true Tm values.',
        'Root Mean Squared Error (RMSE): Square root of MSE, interpretable in the same units as Tm (\u00b0C).',
        'Mean Absolute Error (MAE): Average absolute difference between predictions and true values.',
        'R\u00b2 Score: Coefficient of determination, measuring the proportion of variance explained by the model.',
        'Pearson Correlation Coefficient (r): Measures the linear correlation between predicted and true Tm values.',
    ]
    for m in metrics:
        p = doc.add_paragraph(m, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT_NAME

    doc.add_page_break()

    # ========== CHAPTER 4: RESULTS AND DISCUSSION ==========
    add_heading_styled(doc, 'CHAPTER 4', level=1)
    add_heading_styled(doc, 'RESULTS AND DISCUSSION', level=1)

    add_heading_styled(doc, '4.1 Training Performance', level=2)
    add_para(doc, (
        'Both the GCN and GIN models were trained on the 100,000-protein dataset. The training dynamics are '
        'shown in the training loss curves below.'
    ))

    add_heading_styled(doc, '4.1.1 GCN Training Curve', level=3)
    add_image_if_exists(doc, os.path.join(PROJECT_DIR, 'training_curve.png'),
                       caption='Figure 1: GCN training and test loss curves over 24 epochs.')

    add_para(doc, (
        'The GCN model showed steady convergence with the training MSE loss decreasing from approximately '
        '151 at epoch 0 to approximately 90 by epoch 24. The test loss stabilized around 110-115, indicating '
        'good generalization with a moderate train-test gap. No severe overfitting was observed.'
    ))

    add_heading_styled(doc, '4.1.2 GIN Training Curve', level=3)
    add_image_if_exists(doc, os.path.join(PROJECT_DIR, 'gin_training_curve.png'),
                       caption='Figure 2: GIN training and test loss curves over 24 epochs.')

    add_para(doc, (
        'The GIN model displayed more volatile training dynamics initially, with training loss starting at '
        'approximately 410 and showing oscillations in the early epochs (notably spikes at epochs 3-5). '
        'However, the model converged rapidly after epoch 8, with the training loss reaching approximately '
        '95 and the test loss stabilizing around 110-120. Early stopping was triggered at epoch 24, with the '
        'best model checkpoint saved from the epoch with lowest test loss.'
    ))

    add_heading_styled(doc, '4.2 Model Comparison', level=2)
    add_para(doc, (
        'The two GNN architectures were compared on the test set using the evaluation metrics. The results '
        'are summarized in the table below.'
    ))

    make_table(doc,
        ['Metric', 'GCN Model', 'GIN Model', 'Improvement'],
        [
            ['Pearson Correlation (r)', '0.720', '0.745', '+0.025'],
            ['Test MSE (best)', '~110', '~108', 'GIN slightly better'],
            ['Architecture', '3\u00d7 GCNConv (64d)', '2\u00d7 GINConv (128d)', 'GIN more expressive'],
            ['Pooling', 'Global Mean', 'Global Sum', 'Sum preserves size info'],
            ['Parameters', '~21.5K', '~126K', 'GIN is 6\u00d7 larger'],
            ['Early Stopping', 'Epoch ~24', 'Epoch ~24', 'Similar convergence'],
        ]
    )
    add_para(doc, 'Table 4: Performance comparison between GCN and GIN models.',
             italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, (
        'The GIN model outperformed the GCN model with a Pearson correlation of 0.745 compared to 0.720. '
        'This improvement is attributed to the more expressive message-passing mechanism of GIN, which uses '
        'learnable MLPs for feature aggregation rather than the fixed spectral convolution of GCN. The larger '
        'hidden dimension (128 vs 64) and sum pooling strategy also contribute to the GIN model\'s ability '
        'to capture fine-grained structural features associated with thermal stability.'
    ))

    add_heading_styled(doc, '4.3 Prediction Scatter Plots', level=2)
    add_heading_styled(doc, '4.3.1 GCN Predictions', level=3)
    add_image_if_exists(doc, os.path.join(PROJECT_DIR, 'prediction_scatter.png'),
                       caption='Figure 3: GCN predicted vs. true melting temperatures (Pearson r = 0.720).')

    add_heading_styled(doc, '4.3.2 GIN Predictions', level=3)
    add_image_if_exists(doc, os.path.join(PROJECT_DIR, 'gin_prediction_scatter.png'),
                       caption='Figure 4: GIN predicted vs. true melting temperatures (Pearson r = 0.745).')

    add_para(doc, (
        'The scatter plots reveal several important patterns. Both models show a strong positive correlation '
        'between predicted and true Tm values, with predictions generally clustering around the line of '
        'perfect prediction (dashed red line). The GIN model shows slightly tighter clustering, particularly '
        'in the high-temperature region (70-100\u00b0C), where accurate prediction is most critical for '
        'thermostable protein engineering.'
    ))
    add_para(doc, (
        'Notable observations from the scatter plots include: (1) Both models predict well in the dense '
        'region (25-50\u00b0C) where most proteins reside; (2) Predictions in the extreme temperature ranges '
        '(<15\u00b0C and >90\u00b0C) show higher variance, likely due to fewer training examples; (3) The '
        'vertical banding pattern visible in both plots reflects the discrete Tm values in the dataset '
        '(many proteins share similar experimental Tm); (4) A few outlier predictions exist (e.g., predictions '
        '>120\u00b0C for proteins with true Tm around 50\u00b0C), suggesting that some protein graphs have '
        'unusual features that confuse the model.'
    ))

    add_heading_styled(doc, '4.4 Discussion', level=2)

    add_heading_styled(doc, '4.4.1 Effectiveness of Structure-Guided GNN Approach', level=3)
    add_para(doc, (
        'The results demonstrate that combining ESM-2 protein language model embeddings with structure-based '
        'graph representations provides a powerful framework for absolute Tm prediction. A Pearson correlation '
        'of 0.745 on a diverse dataset of 100,000 proteins is a strong result, particularly considering the '
        'inherent noise in experimental Tm measurements (which can vary by 2-5\u00b0C across different studies '
        'and experimental conditions) and the diversity of protein families in the dataset.'
    ))
    add_para(doc, (
        'The structure-guided approach is effective because the 8\u00c5 contact graph captures long-range '
        'interactions (salt bridges, hydrophobic cores, disulfide bonds) that stabilize the protein fold but '
        'are not directly accessible from sequence alone. The AlphaFold-predicted structures, while not '
        'experimental, provide sufficiently accurate spatial information for the GNN to learn meaningful '
        'structural patterns associated with thermal stability.'
    ))

    add_heading_styled(doc, '4.4.2 GIN vs GCN: Architectural Insights', level=3)
    add_para(doc, (
        'The superior performance of GIN over GCN can be attributed to several factors. First, GIN\'s learnable '
        'MLP-based aggregation is strictly more expressive than GCN\'s spectral convolution, allowing it to '
        'distinguish a wider class of graph structures. This is theoretically grounded in the Weisfeiler-Lehman '
        'graph isomorphism framework. Second, the larger hidden dimension (128 vs 64) provides more capacity '
        'to encode the rich 320-dimensional ESM-2 embeddings. Third, global sum pooling preserves information '
        'about protein size (number of residues), which may correlate with molecular weight and thus stability.'
    ))

    add_heading_styled(doc, '4.4.3 Comparison with Preliminary Results', level=3)
    add_para(doc, (
        'During the development phase, an initial GCN model was trained on a 500-sample development subset, '
        'achieving a test MSE of approximately 101 (corresponding to an average prediction error of \u224810\u00b0C). '
        'Scaling to the full 100,000-sample dataset improved the Pearson correlation from a development-phase '
        'baseline to 0.720 (GCN) and 0.745 (GIN), demonstrating the benefit of larger, more diverse training '
        'data for learning generalizable stability patterns.'
    ))

    add_heading_styled(doc, '4.4.4 Limitations', level=3)
    add_para(doc, 'Several limitations of the current study should be noted:')
    limitations = [
        'The dataset contains multiple entries for the same protein under different experimental conditions, which may introduce data leakage between train and test sets.',
        'AlphaFold-predicted structures may deviate from experimental structures, particularly for disordered regions and multi-domain proteins.',
        'The 8\u00c5 threshold for edge construction is a single fixed choice; adaptive or multi-scale thresholds may improve performance.',
        'Sequences longer than 1,022 residues are truncated for ESM-2 embedding extraction, potentially losing information for very large proteins.',
        'No external validation on completely independent datasets was performed.',
        'The model does not account for experimental conditions (pH, buffer, ionic strength) that influence measured Tm values.',
    ]
    for lim in limitations:
        p = doc.add_paragraph(lim, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT_NAME

    add_heading_styled(doc, '4.4.5 Future Directions', level=3)
    add_para(doc, 'Several promising directions for future work include:')
    future = [
        'Multi-scale graph construction: Using multiple distance thresholds (4\u00c5, 8\u00c5, 12\u00c5) to capture contacts at different spatial scales.',
        'Attention-based GNNs: Implementing Graph Attention Networks (GAT) or Graph Transformers to learn which contacts are most important for stability.',
        'Multi-modal fusion: Combining graph-level features with global descriptors (molecular weight, isoelectric point, amino acid composition) for hybrid prediction.',
        'Uncertainty quantification: Adding dropout-based or ensemble-based uncertainty estimates to predictions, enabling confidence-weighted screening.',
        'Cross-family generalization: Evaluating model performance on held-out protein families to assess transferability.',
        'Fine-grained ESM-2 models: Using larger ESM-2 variants (e.g., esm2_t33_650M) for richer residue embeddings.',
        'Active learning: Prioritizing proteins for experimental Tm validation based on model uncertainty.',
    ]
    for f in future:
        p = doc.add_paragraph(f, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT_NAME

    doc.add_page_break()

    # ========== CHAPTER 5: CONCLUSION ==========
    add_heading_styled(doc, 'CHAPTER 5', level=1)
    add_heading_styled(doc, 'CONCLUSION', level=1)
    add_para(doc, (
        'This study successfully developed and evaluated a structure-guided Graph Neural Network system for '
        'predicting protein melting temperature from amino acid sequences and three-dimensional structural '
        'information. The complete pipeline\u2014from raw FASTA data ingestion, through AlphaFold structure '
        'retrieval, ESM-2 embedding extraction, graph construction, and GNN-based regression\u2014was implemented '
        'and validated on a large-scale dataset of 100,000 proteins from the TemStaPro database.'
    ))
    add_para(doc, (
        'Two GNN architectures were systematically compared. The Graph Isomorphism Network (GIN) achieved '
        'a Pearson correlation coefficient of 0.745, outperforming the Graph Convolutional Network (GCN) '
        'which achieved r = 0.720. This improvement is attributed to GIN\'s more expressive message-passing '
        'mechanism, larger model capacity, and sum-based graph pooling. Both models demonstrated stable '
        'training dynamics with convergence around epoch 24 and no severe overfitting, benefiting from '
        'L2 regularization and early stopping.'
    ))
    add_para(doc, (
        'The results demonstrate that combining protein language model embeddings (ESM-2) with structure-based '
        'residue contact graphs provides a powerful and scalable framework for absolute Tm prediction. Unlike '
        'sequence-only approaches that ignore spatial interactions, and unlike \u0394\u0394G-based methods '
        'that require a reference protein, our approach directly predicts absolute melting temperatures for '
        'arbitrary protein sequences with available or predicted 3D structures.'
    ))
    add_para(doc, (
        'The pipeline is designed for scalability and extensibility. The modular architecture allows easy '
        'substitution of embedding models, graph construction strategies, and GNN architectures. Future work '
        'will focus on incorporating attention mechanisms, multi-scale graph construction, and experimental '
        'condition features to further improve prediction accuracy. Overall, this work establishes a strong '
        'foundation for structure-aware protein stability prediction using graph neural networks, with '
        'potential applications in therapeutic protein engineering, industrial enzyme design, and computational '
        'biology research.'
    ))

    doc.add_page_break()

    # ========== REFERENCES ==========
    add_heading_styled(doc, 'REFERENCES', level=1)

    refs = [
        '[1] Ku, T., Lu, P., Chan, C., Wang, T., Lai, S., Lyu, P., & Hsiao, N. (2009). Predicting melting temperature directly from protein sequences. Computational Biology and Chemistry, 33(6), 445-450.',
        '[2] Pucci, F., Rooman, M. (2016). Improved metrics for protein thermostability prediction. BMC Bioinformatics.',
        '[3] Lin, Z., Akin, H., Rao, R., Hie, B., Zhu, Z., Lu, W., ... & Rives, A. (2023). Evolutionary-scale prediction of atomic-level protein structure with a language model. Science, 379(6637), 1123-1130.',
        '[4] Jumper, J., Evans, R., Pritzel, A., Green, T., Figurnov, M., Ronneberger, O., ... & Hassabis, D. (2021). Highly accurate protein structure prediction with AlphaFold. Nature, 596(7873), 583-589.',
        '[5] Kipf, T. N., & Welling, M. (2017). Semi-supervised classification with graph convolutional networks. arXiv preprint arXiv:1609.02907.',
        '[6] Xu, K., Hu, W., Leskovec, J., & Jegelka, S. (2019). How powerful are graph neural networks? arXiv preprint arXiv:1810.00826.',
        '[7] Elnaggar, A., Heinzinger, M., Dallago, C., Rehber, G., Wang, Y., Jones, L., ... & Rost, B. (2022). ProtTrans: Toward Understanding the Language of Life Through Self-Supervised Learning. IEEE Transactions on Pattern Analysis and Machine Intelligence, 44(10).',
        '[8] ThermoAGT\u2011GA (2024). Predicting protein thermal stability changes upon single and multi-point mutations via restricted attention subgraph neural network. Journal of Theoretical Biology.',
        '[9] ProSTAGE (2024). Predicting Effects of Mutations on Protein Stability by Fusing Structure and Sequence Embedding. Journal of Chemical Information and Modeling.',
        '[10] Stability Oracle (2024). Structure-based graph-transformer framework for identifying thermodynamically stable protein sequences. Nature Communications.',
        '[11] TemBERTure (2024). Advancing protein thermostability prediction with deep learning. Bioinformatics Advances.',
        '[12] DeepTM (2023). Deep learning algorithm for prediction of melting temperature of thermophilic proteins directly from sequences. Computational and Structural Biotechnology Journal.',
        '[13] TemStaPro (2024). TemStaPro: protein thermostability prediction using sequence representations from protein language models. Bioinformatics, 40(4).',
        '[14] Fey, M., & Lenssen, J. E. (2019). Fast graph representation learning with PyTorch Geometric. arXiv preprint arXiv:1903.02428.',
        '[15] Rives, A., Meier, J., Sercu, T., Goyal, S., Lin, Z., Liu, J., ... & Fergus, R. (2021). Biological structure and function emerge from scaling unsupervised learning to 250 million protein sequences. Proceedings of the National Academy of Sciences, 118(15).',
        '[16] Scarselli, F., Gori, M., Tsoi, A. C., Hagenbuchner, M., & Monfardini, G. (2009). The graph neural network model. IEEE Transactions on Neural Networks, 20(1), 61-80.',
        '[17] Kingma, D. P., & Ba, J. (2015). Adam: A method for stochastic optimization. arXiv preprint arXiv:1412.6980.',
        '[18] Srivastava, N., Hinton, G., Krizhevsky, A., Sutskever, I., & Salakhutdinov, R. (2014). Dropout: a simple way to prevent neural networks from overfitting. JMLR, 15(1), 1929-1958.',
        '[19] Ioffe, S., & Szegedy, C. (2015). Batch normalization: Accelerating deep network training by reducing internal covariate shift. ICML, 448-456.',
        '[20] Varadi, M., Anyango, S., Deshpande, M., Nair, S., Natassia, C., Yordanova, G., ... & Velankar, S. (2022). AlphaFold Protein Structure Database: massively expanding the structural coverage of protein-sequence space with high-accuracy models. Nucleic Acids Research, 50(D1), D439-D444.',
    ]

    for ref in refs:
        add_para(doc, ref, size=11, space_after=4)

    # ========== SAVE ==========
    output_path = os.path.join(PROJECT_DIR, 'Team-22_IBS_REPORT_Achintya_TR_v2.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')

if __name__ == '__main__':
    generate_report()
